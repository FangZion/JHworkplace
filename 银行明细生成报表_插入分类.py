"""
集团资金月度流量分析报告 — 自动生成（可定制分类版）
=========================================================
读取「明细合并」，基于内转标记/货款收入标记分类，
所有统计值按币种独立核算，不跨币种折算。

收入类型：客户货款 | 利息/理财 | 政府补贴/退税 | 光伏收入 | 房租水电 | 贷款/融资/贴现 | 集团内收付 | 其他收入
支出类型：供应商货款 | 薪资/社保/公积金 | 银行本利/手续费 | 承兑保证金 | 基建工程款 | 集团内收付 | 税款 | 货代费 | 其他支出

使用：python generate_report.py
"""

import argparse, calendar, datetime, io, os, sys
from collections import Counter, defaultdict
from dataclasses import dataclass
from typing import Dict, List, Optional, Tuple

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import numpy as np

import openpyxl
from openpyxl.reader.excel import load_workbook as _load_workbook
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor, Inches

# matplotlib 中文字体
_fp = r"C:\Windows\Fonts\msyh.ttc"
_fp_prop = fm.FontProperties(fname=_fp)
plt.rcParams["font.family"] = _fp_prop.get_name()
fm.fontManager.addfont(_fp)
plt.rcParams["axes.unicode_minus"] = False
CLR_PRIMARY = "#C00000"
CLR_ACCENT  = "#4472C4"
CLR_DARK    = "#333333"


# ============================================================
# 0. 修复 XLSX 样式兼容性
# ============================================================
def _fix_xlsx_styles(src_path: str) -> str:
    """
    修复 openpyxl 无法解析的空白 Fill 样式。
    返回修复后的文件路径（覆盖原文件），如果无需修复则返回原路径。
    """
    import zipfile, tempfile
    from xml.etree import ElementTree as ET

    NS = 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'
    ET.register_namespace('', NS)

    fixed_path = None
    with zipfile.ZipFile(src_path, 'r') as zin:
        data = {}
        needs_fix = False
        for name in zin.namelist():
            content = zin.read(name)
            data[name] = content
            if name == 'xl/styles.xml':
                root = ET.fromstring(content)
                fills_elem = root.find(f'{{{NS}}}fills')
                if fills_elem is None:
                    continue
                for fill in fills_elem:
                    if len(fill) == 0:
                        needs_fix = True
                        # 替换为空 patternFill
                        pattern = ET.SubElement(fill, f'{{{NS}}}patternFill')
                        pattern.set('patternType', 'none')
                if needs_fix:
                    data[name] = ET.tostring(root, xml_declaration=True, encoding='UTF-8')

        if needs_fix:
            fd, fixed_path = tempfile.mkstemp(suffix='.xlsx')
            os.close(fd)
            with zipfile.ZipFile(fixed_path, 'w', zipfile.ZIP_DEFLATED) as zout:
                for name, content in data.items():
                    zout.writestr(name, content)
            return fixed_path
    return src_path


def _load_workbook_safe(filepath: str, data_only: bool = True):
    """加载工作簿，自动修复样式兼容性问题"""
    safe_path = _fix_xlsx_styles(filepath)
    try:
        return _load_workbook(safe_path, data_only=data_only)
    finally:
        if safe_path != filepath and os.path.exists(safe_path):
            try:
                os.remove(safe_path)
            except OSError:
                pass

# ============================================================
# 1. 分类映射 — 严格按R列（第18列）归集
# ============================================================
# 全量分类列表（收入+支出）
ALL_CATEGORIES = [
    "客户货款", "利息/理财", "政府补贴/退税", "光伏收入", "房租水电",
    "贷款/融资/贴现", "集团内调拨", "集团内收付",
    "供应商货款", "薪资/社保/公积金", "银行本利/手续费", "承兑保证金",
    "基建工程款", "税款", "货代费",
]

# 分类 → 现金流量活动类型
CATEGORY_ACTIVITY = {
    "客户货款": "经营活动",
    "利息/理财": "投资活动",
    "政府补贴/退税": "经营活动",
    "光伏收入": "经营活动",
    "房租水电": "经营活动",
    "贷款/融资/贴现": "筹资活动",
    "集团内调拨": "经营活动",
    "集团内收付": "经营活动",
    "供应商货款": "经营活动",
    "薪资/社保/公积金": "经营活动",
    "银行本利/手续费": "经营活动",
    "承兑保证金": "经营活动",
    "基建工程款": "投资活动",
    "税款": "经营活动",
    "货代费": "经营活动",
}

# 收入端分类列表（用于报表展示顺序）
INCOME_CATS_ORDER = [
    "客户货款", "利息/理财", "政府补贴/退税", "光伏收入", "房租水电",
    "贷款/融资/贴现", "集团内收付",
]

# 支出端分类列表（用于报表展示顺序）
EXPENSE_CATS_ORDER = [
    "供应商货款", "薪资/社保/公积金", "银行本利/手续费", "承兑保证金",
    "基建工程款", "集团内收付", "税款", "货代费",
]

INCOME_FALLBACK = "其他收入"
EXPENSE_FALLBACK = "其他支出"



def classify(tx, is_income: bool) -> Tuple[str, str]:
    """
    返回 (大类, 子类)
    大类 = 经营活动/投资活动/筹资活动/其他
    子类 = 直接取 R列（第18列）的值，严格匹配用户指定分类
    """
    # R列有值且可识别时优先使用（用户指定分类，不做关键字猜测）
    r_cat = tx.r_category.strip()
    if r_cat in CATEGORY_ACTIVITY:
        return (CATEGORY_ACTIVITY[r_cat], r_cat)

    # R列无法匹配时，内转标记作兜底
    if tx.transfer_flag == 1:
        return ("经营活动", "集团内收付")

    # R列为空或无法匹配时的兜底
    if is_income:
        if tx.income_flag == 1:
            return ("经营活动", "客户货款")
        return ("其他", INCOME_FALLBACK)
    else:
        return ("其他", EXPENSE_FALLBACK)


# ============================================================
# 2. 数据模型
# ============================================================
@dataclass
class Transaction:
    date: datetime.datetime
    note: str; currency: str; income: float; expense: float
    company: str; counterparty: str
    transfer_flag: int; income_flag: int
    # 新增字段
    trade_no: str = ""          # 交易流水号 (B列)
    balance: float = 0.0         # 账户余额 (G列)
    bank_self: str = ""         # 本方开户行 (I列)
    account_self: str = ""      # 本方账号 (J列)
    bank_counter: str = ""      # 对方开户行 (L列)
    account_counter: str = ""   # 对方账号 (M列)
    trade_note: str = ""        # 交易备注 (O列)
    r_category: str = ""        # R列分类


FOREX_KW = ["换汇", "网银结汇", "结汇", "购汇"]

def is_forex(note: str) -> bool:
    return any(k in note for k in FOREX_KW)


def parse_excel_date(raw) -> Optional[datetime.datetime]:
    if raw is None: return None
    if isinstance(raw, datetime.datetime): return raw
    if isinstance(raw, (int, float)):
        try: return datetime.datetime.fromordinal(datetime.datetime(1900,1,1).toordinal()+int(raw)-2)
        except: return None
    return None


def get_base_dir() -> str:
    if getattr(sys, "frozen", False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))


def parse_month_from_filename(filename: str) -> Optional[str]:
    """从文件名提取月份，如 资金台账2026.05.xlsx → 2026-05"""
    import re
    m = re.search(r'(\d{4})\.(\d{2})', filename)
    if m:
        return f"{m.group(1)}-{m.group(2)}"
    return None


def scan_input_files(base_dir: str) -> List[Tuple[str, str]]:
    """扫描目录，返回 [(文件路径, 月份), ...] 按月份降序"""
    files = []
    for f in sorted(os.listdir(base_dir)):
        if not f.endswith(".xlsx") or f.startswith("~$") or f.startswith("资金明细_"):
            continue
        month = parse_month_from_filename(f)
        if month:
            files.append((os.path.join(base_dir, f), month))
    files.sort(key=lambda x: x[1], reverse=True)
    return files


def select_file_interactive(files: List[Tuple[str, str]]) -> Tuple[str, str]:
    """方向键选择台账文件（↑/↓ 移动，Enter 确认），返回 (文件路径, 月份)"""
    import msvcrt

    idx = 0
    while True:
        import os as _os
        _os.system("cls")
        print("=" * 52)
        print("   📅 请选择要生成报告的台账")
        print("=" * 52)
        print("   ↑/↓ 移动光标  Enter 确认\n")
        for i, (fp, m) in enumerate(files):
            fname = os.path.basename(fp)
            prefix = "  ▸ " if i == idx else "    "
            print(f"{prefix}{m}  {fname}")
        print()

        key = msvcrt.getch()
        if key == b'\xe0':          # 方向键前缀
            key2 = msvcrt.getch()
            if key2 == b'H':         # ↑
                idx = (idx - 1) % len(files)
            elif key2 == b'P':       # ↓
                idx = (idx + 1) % len(files)
        elif key == b'\r':           # Enter
            _os.system("cls")
            fp, m = files[idx]
            print(f"✅ 已选择：{m}  {os.path.basename(fp)}\n")
            return (fp, m)
        elif key == b'\x1b':         # Esc
            _os.system("cls")
            print("已取消\n")
            sys.exit(0)


def load_exchange_rates(filepath: str) -> dict:
    """从资金日报读取汇率，返回 {币种: 兑人民币汇率}"""
    wb = _load_workbook_safe(filepath, data_only=True)
    ws = wb["资金日报"]
    # Row5: [今日汇率, None, 人民币汇率, 美元汇率, 欧元汇率, 泰铢汇率, 港币汇率, 日元汇率, ...]
    # Col3=人民币, Col4=美元, Col5=欧元, Col6=泰铢, Col7=港币, Col8=日元
    ccy_col = {"人民币":3, "美元":4, "欧元":5, "泰铢":6, "港币":7, "日元":8}
    rates = {}
    for ccy, col in ccy_col.items():
        val = ws.cell(5, col).value
        if val: rates[ccy] = float(val)
    rates["人民币"] = 1.0
    return rates


def load_transactions(filepath: str, month_str: str) -> List[Transaction]:
    wb = _load_workbook_safe(filepath, data_only=True)
    ws = wb["明细合并"]
    ym = datetime.datetime.strptime(month_str, "%Y-%m")
    start = ym.replace(day=1)
    end = ym.replace(year=ym.year+1, month=1, day=1) if ym.month==12 else ym.replace(month=ym.month+1, day=1)

    txns = []
    for r in range(2, ws.max_row+1):
        dt = parse_excel_date(ws.cell(r,1).value)
        if dt is None or dt < start or dt >= end: continue
        t = Transaction(
            date=dt,
            note=str(ws.cell(r,14).value or "").strip(),
            currency=str(ws.cell(r,4).value or "人民币"),
            income=float(ws.cell(r,5).value or 0),
            expense=float(ws.cell(r,6).value or 0),
            company=str(ws.cell(r,8).value or ""),
            counterparty=str(ws.cell(r,11).value or ""),
            transfer_flag=1 if ws.cell(r,32).value==1 else 0,
            income_flag=1 if ws.cell(r,33).value==1 else 0,
            trade_no=str(ws.cell(r,2).value or "").strip(),
            balance=float(ws.cell(r,7).value or 0),
            bank_self=str(ws.cell(r,9).value or "").strip(),
            account_self=str(ws.cell(r,10).value or "").strip(),
            bank_counter=str(ws.cell(r,12).value or "").strip(),
            account_counter=str(ws.cell(r,13).value or "").strip(),
            trade_note=str(ws.cell(r,15).value or "").strip(),
            r_category=str(ws.cell(r,18).value or "").strip(),
        )
        if t.income==0 and t.expense==0: continue
        txns.append(t)
    txns.sort(key=lambda x: x.date)
    return txns


# ============================================================
# 4. 分析
# ============================================================
def analyze(txns: List[Transaction], exchange_rates: dict = None) -> dict:
    by_ccy: dict = defaultdict(lambda: defaultdict(lambda: {"income":0.0,"expense":0.0,"count":0}))
    for t in txns:
        c = t.currency
        if t.income>0:
            a, s = classify(t, True)
            by_ccy[c][a]["income"]+=t.income; by_ccy[c][a]["count"]+=1
        if t.expense>0:
            a, s = classify(t, False)
            by_ccy[c][a]["expense"]+=t.expense; by_ccy[c][a]["count"]+=1

    out = {c: dict(v) for c,v in by_ccy.items()}
    ccy_order = sorted(out, key=lambda c: -sum(v["income"]+v["expense"] for v in out[c].values()))
    er = exchange_rates or {}

    # === TOP 交易 ===
    # 收入 TOP：排除集团内收付、贷款融资（含贷款/融资/贴现）
    top_in = []
    for t in txns:
        if t.income<=0: continue
        a,s=classify(t,True)
        if s in ("集团内调拨", "集团内收付", "贷款/融资/贴现"): continue
        top_in.append(t)
    top_in = sorted(top_in, key=lambda t: -to_rmb_(t.income,t.currency,er))[:10]

    # 支出 TOP：仅统计供应商货款、基建工程款、税款、货代费、其他支出五类
    EXP_TOP_CATS = {"供应商货款", "基建工程款", "税款", "货代费", "其他支出"}
    top_ex = []
    for t in txns:
        if t.expense<=0: continue
        a,s=classify(t,False)
        if s not in EXP_TOP_CATS: continue
        top_ex.append(t)
    top_ex = sorted(top_ex, key=lambda t: -to_rmb_(t.expense,t.currency,er))[:10]

    # 公司排名（不含内部调拨）
    co_by_ccy: dict = defaultdict(lambda: defaultdict(lambda: {"income":0.0,"expense":0.0,"count":0}))
    for t in txns:
        if t.transfer_flag==1: continue
        co_by_ccy[t.currency][t.company]["income"]+=t.income
        co_by_ccy[t.currency][t.company]["expense"]+=t.expense
        co_by_ccy[t.currency][t.company]["count"]+=1

    # 每日趋势
    daily: dict = defaultdict(lambda: defaultdict(lambda: {"income":0.0,"expense":0.0,"count":0}))
    for t in txns:
        d = t.date.strftime("%m-%d")
        daily[t.currency][d]["income"]+=t.income
        daily[t.currency][d]["expense"]+=t.expense
        daily[t.currency][d]["count"]+=1

    # 合计
    totals = {}
    for ccy, cats in out.items():
        ti = sum(v["income"] for v in cats.values())
        te = sum(v["expense"] for v in cats.values())
        totals[ccy] = {"income": ti, "expense": te, "net": ti-te}

    # 子类明细
    subcats: dict = defaultdict(lambda: defaultdict(lambda: {"income":0.0,"expense":0.0,"count":0}))
    for t in txns:
        c = t.currency
        if t.income>0:
            a, s = classify(t, True)
            subcats[c][s]["income"]+=t.income; subcats[c][s]["count"]+=1
        if t.expense>0:
            a, s = classify(t, False)
            subcats[c][s]["expense"]+=t.expense; subcats[c][s]["count"]+=1

    # 全币种折人民币汇总
    rmb_sum = defaultdict(lambda: {"income":0.0,"expense":0.0,"count":0})
    for ccy, cats in subcats.items():
        rate = er.get(ccy, 1.0)
        for sub, v in cats.items():
            rmb_sum[sub]["income"] += v["income"] * rate
            rmb_sum[sub]["expense"] += v["expense"] * rate
            rmb_sum[sub]["count"] += v["count"]
    rmb_total_inc = sum(v["income"] for v in rmb_sum.values())
    rmb_total_exp = sum(v["expense"] for v in rmb_sum.values())

    # 分币种收入/支出笔数
    ccy_inc_count = Counter()
    ccy_exp_count = Counter()
    for t in txns:
        if t.income > 0: ccy_inc_count[t.currency] += 1
        if t.expense > 0: ccy_exp_count[t.currency] += 1

    # 注：to_rmb_ 提前引用，在函数外定义
    return {
        "by_ccy": out, "ccy_order": ccy_order, "totals": totals,
        "top_incomes": top_in, "top_expenses": top_ex,
        "company_by_ccy": {k:dict(v) for k,v in co_by_ccy.items()},
        "daily_by_ccy": {k:dict(v) for k,v in daily.items()},
        "subcats": {k:dict(v) for k,v in subcats.items()},
        "rmb_sum": dict(rmb_sum),
        "rmb_total_inc": rmb_total_inc,
        "rmb_total_exp": rmb_total_exp,
        "ccy_inc_count": dict(ccy_inc_count),
        "ccy_exp_count": dict(ccy_exp_count),
        "exchange_rates": er,
        "total_count": len(txns),
    }


# ============================================================
# 5. 报告生成
# ============================================================
def line_chart_bytes(dates, incomes, expenses, title, ccy):
    """生成折线图 PNG，透明背景"""
    fig, ax = plt.subplots(figsize=(6.5, 3.2))
    fig.patch.set_alpha(0); ax.set_facecolor("none")
    x = np.arange(len(dates))
    ax.plot(x, incomes, color=CLR_PRIMARY, marker="o", lw=2, ms=4, label="收入")
    ax.plot(x, expenses, color=CLR_ACCENT, marker="s", lw=2, ms=4, label="支出")
    ax.fill_between(x, incomes, expenses,
        where=np.array(incomes)>=np.array(expenses), interpolate=True, alpha=0.12, color=CLR_PRIMARY)
    ax.fill_between(x, incomes, expenses,
        where=np.array(incomes)<np.array(expenses), interpolate=True, alpha=0.12, color=CLR_ACCENT)
    step = max(1, len(dates)//10)
    ax.set_xticks(x)
    ax.set_xticklabels([d if i%step==0 else "" for i,d in enumerate(dates)], fontsize=7, rotation=30, ha="right")
    ax.legend(loc="upper left", fontsize=9, prop=_fp_prop)
    ax.set_title(title, fontsize=11, fontweight="bold", color=CLR_DARK, pad=10, fontproperties=_fp_prop)
    ax.set_ylabel(ccy, fontsize=9, fontproperties=_fp_prop)
    ax.grid(True, alpha=0.3)
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format="png", dpi=150, bbox_inches="tight", transparent=True)
    plt.close(); buf.seek(0)
    return buf


def to_rmb_(v: float, ccy: str, rates: dict) -> float:
    return v * rates.get(ccy, 1.0)


def fmt(v: float, ccy: str) -> str:
    if ccy in ("泰铢","日元"): return f"{v:,.0f}"
    return f"{v:,.2f}"

def pct(a: float, b: float) -> str:
    return f"{a/max(b,1)*100:.1f}%"

def set_cell(cell, text, bold=False, size=7, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text=""
    p=cell.paragraphs[0]; p.alignment=align
    run=p.add_run(str(text)); run.font.size=Pt(size); run.bold=bold
    run.font.name="微软雅黑"; run._element.rPr.rFonts.set(qn("w:eastAsia"),"微软雅黑")

def make_table(doc, headers, rows, cw=None):
    t=doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
    for ci,h in enumerate(headers): set_cell(t.rows[0].cells[ci],h,bold=True,size=7)
    for ri,r in enumerate(rows):
        for ci,v in enumerate(r): set_cell(t.rows[ri+1].cells[ci],v,bold=(ri==len(rows)-1),size=7)
    if cw:
        for row in t.rows:
            for ci,w in enumerate(cw):
                if ci<len(row.cells): row.cells[ci].width=Cm(w)
    return t

def add_h(doc, text, level=2):
    h=doc.add_heading(text,level=level)
    for run in h.runs:
        run.font.name="微软雅黑"; run._element.rPr.rFonts.set(qn("w:eastAsia"),"微软雅黑")


def build_report(txns, month_str, output_path, stats):
    doc=Document()
    style=doc.styles["Normal"]; style.font.name="微软雅黑"; style.font.size=Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"),"微软雅黑")

    ym=datetime.datetime.strptime(month_str,"%Y-%m")
    ml=f"{ym.year}年{ym.month}月"
    ld=calendar.monthrange(ym.year,ym.month)[1]

    # 标题
    title=doc.add_heading("",level=0); title.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=title.add_run("集团资金月度流量分析报告")
    run.font.size=Pt(20); run.bold=True; run.font.name="微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"),"微软雅黑")
    sub=doc.add_paragraph(); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=sub.add_run(f"报告期间：{ml}1日 — {ml}{ld}日")
    r.font.size=Pt(12); r.font.color.rgb=RGBColor(100,100,100)
    doc.add_paragraph()

    ccy_order=stats["ccy_order"]; by_ccy=stats["by_ccy"]; totals=stats["totals"]

    # === 一、资金全景概览 ===
    add_h(doc,"一、资金全景概览",level=1)
    inc_cnt=stats.get("ccy_inc_count",{})
    exp_cnt=stats.get("ccy_exp_count",{})
    overview=[]
    for ccy in ccy_order:
        t2=totals[ccy]
        nl=f"+{fmt(t2['net'],ccy)}" if t2['net']>=0 else fmt(t2['net'],ccy)
        overview.append([ccy, fmt(t2['income'],ccy), fmt(t2['expense'],ccy), nl,
                         str(inc_cnt.get(ccy,0)), str(exp_cnt.get(ccy,0))])
    make_table(doc,["币种","总收入","总支出","净流量","收入笔数","支出笔数"],overview, cw=[2.5,3.5,3.5,3,2.5,2.5])
    p=doc.add_paragraph()
    r=p.add_run(f"共 {stats['total_count']} 笔交易，{len(ccy_order)} 种币种。各币种独立核算，未跨币种折算。")
    r.font.size=Pt(10); r.font.color.rgb=RGBColor(120,120,120)
    doc.add_paragraph()

    # === 二、流向分析 ===
    add_h(doc,"二、资金流向分析",level=1)

    inc_order = ["客户货款","利息/理财","政府补贴/退税","光伏收入","房租水电","贷款/融资/贴现","集团内收付","其他收入"]
    exp_order = ["供应商货款","薪资/社保/公积金","银行本利/手续费","承兑保证金","基建工程款","集团内收付","税款","货代费","其他支出"]

    # ---- 全币种折人民币汇总 ----
    add_h(doc,"▸ 全币种折人民币汇总",level=2)
    rmb=stats.get("rmb_sum",{}); rmb_inc=stats.get("rmb_total_inc",0); rmb_exp=stats.get("rmb_total_exp",0)
    irows=[]
    for sub in inc_order:
        v=rmb.get(sub,{}); amt=v.get("income",0)
        if amt>0: irows.append([sub,f"{amt/10000:,.2f}",pct(amt,rmb_inc)])
    irows.append(["合计",f"{rmb_inc/10000:,.2f}","100%"])
    p=doc.add_paragraph("收入结构（万元人民币）：")
    make_table(doc,["收入类型","金额(万元)","占比"],irows,cw=[5,6,3])
    doc.add_paragraph()
    erows=[]
    for sub in exp_order:
        v=rmb.get(sub,{}); amt=v.get("expense",0)
        if amt>0: erows.append([sub,f"{amt/10000:,.2f}",pct(amt,rmb_exp)])
    erows.append(["合计",f"{rmb_exp/10000:,.2f}","100%"])
    p=doc.add_paragraph("支出结构（万元人民币）：")
    make_table(doc,["支出类型","金额(万元)","占比"],erows,cw=[5,6,3])
    doc.add_paragraph()
    rates=stats.get("exchange_rates",{})
    if rates:
        rs="  |  ".join(f"1 {k}={v}" for k,v in sorted(rates.items()) if k!="人民币")
        p=doc.add_paragraph(); r=p.add_run(f"汇率参考（1外币→人民币）：{rs}")
        r.font.size=Pt(8); r.font.color.rgb=RGBColor(0x99,0x99,0x99)
        r.font.name="微软雅黑"; r._element.rPr.rFonts.set(qn("w:eastAsia"),"微软雅黑")
    doc.add_paragraph()

    # === 三、大额交易 TOP10（含折人民币，按折人民币排序） ===
    add_h(doc,"三、大额交易 TOP10（不含换汇、内转）",level=1)
    rates=stats.get("exchange_rates",{})

    add_h(doc,"收入 TOP10（按折人民币）",level=2)
    inc_rmb=sorted([(t,to_rmb_(t.income,t.currency,rates)) for t in stats["top_incomes"]],key=lambda x:-x[1])[:10]
    it=[]
    for t,rmb in inc_rmb:
        a,s=classify(t,True)
        it.append([t.date.strftime("%m-%d"),t.company[:12],t.counterparty[:22],
                   fmt(t.income,t.currency),t.currency,f"{rmb/10000:,.2f}万",s])
    make_table(doc,["日期","公司","对方户名","金额","币种","折RMB","分类"],it,cw=[2,3,4.5,2.5,1.5,2.5,2])
    doc.add_paragraph()

    add_h(doc,"支出 TOP10（按折人民币）",level=2)
    exp_rmb=sorted([(t,to_rmb_(t.expense,t.currency,rates)) for t in stats["top_expenses"]],key=lambda x:-x[1])[:10]
    et=[]
    for t,rmb in exp_rmb:
        a,s=classify(t,False)
        et.append([t.date.strftime("%m-%d"),t.company[:12],t.counterparty[:22],
                   fmt(t.expense,t.currency),t.currency,f"{rmb/10000:,.2f}万",s])
    make_table(doc,["日期","公司","对方户名","金额","币种","折RMB","分类"],et,cw=[2,3,4.5,2.5,1.5,2.5,2])
    doc.add_paragraph()

    # === 四、每日趋势（折线图） ===
    add_h(doc,"四、每日资金流量趋势",level=1)
    rates=stats.get("exchange_rates",{})

    # 全币种折人民币每日汇总
    rmb_daily=defaultdict(lambda:{"income":0.0,"expense":0.0})
    for ccy in ccy_order:
        dd=stats["daily_by_ccy"].get(ccy,{})
        if not dd: continue
        rate=rates.get(ccy,1.0)
        for d,v in dd.items():
            rmb_daily[d]["income"]+=v["income"]*rate
            rmb_daily[d]["expense"]+=v["expense"]*rate
    if rmb_daily:
        days=sorted(rmb_daily.keys())
        inc=[rmb_daily[d]["income"]/10000 for d in days]
        exp=[rmb_daily[d]["expense"]/10000 for d in days]
        buf=line_chart_bytes(days,inc,exp,"全币种折人民币 每日收支趋势（万元）","万元")
        p=doc.add_paragraph(); r=p.add_run("全币种折人民币 每日收支趋势")
        r.bold=True; r.font.size=Pt(12); r.font.name="微软雅黑"
        r._element.rPr.rFonts.set(qn("w:eastAsia"),"微软雅黑")
        doc.add_picture(buf,width=Inches(6))
        doc.add_paragraph()

    for ccy in ccy_order:
        dd=stats["daily_by_ccy"].get(ccy,{})
        if not dd: continue
        days=sorted(dd.keys())
        inc=[dd[d]["income"] for d in days]
        exp=[dd[d]["expense"] for d in days]
        p=doc.add_paragraph(); r=p.add_run(f"{ccy} 每日收支趋势")
        r.bold=True; r.font.size=Pt(12); r.font.name="微软雅黑"
        r._element.rPr.rFonts.set(qn("w:eastAsia"),"微软雅黑")
        buf=line_chart_bytes(days,inc,exp,f"",ccy)
        doc.add_picture(buf,width=Inches(6))
        doc.add_paragraph()

    # === 五、工作量统计（不含手续费） ===
    add_h(doc,"五、工作量统计",level=1)
    p=doc.add_paragraph('支出记录（已过滤备注含"手续费"的交易）：')
    # 按公司统计支出笔数 + 折人民币金额分档
    work=defaultdict(lambda:{"total":0,"small":0,"large":0})
    for t in txns:
        if t.expense<=0: continue
        if "手续费" in t.note: continue
        rmb=to_rmb_(t.expense,t.currency,rates)
        c=t.company or "未知"
        work[c]["total"]+=1
        if rmb<=3000: work[c]["small"]+=1
        else: work[c]["large"]+=1
    ws_rows=[[co[:18],str(v["total"]),str(v["small"]),str(v["large"])] for co,v in sorted(work.items(),key=lambda x:-x[1]["total"])]
    tot_small=sum(v["small"] for v in work.values())
    tot_large=sum(v["large"] for v in work.values())
    ws_rows.append(["合计",str(sum(v["total"] for v in work.values())),str(tot_small),str(tot_large)])
    make_table(doc,["公司","总笔数","≤3000元",">3000元"],ws_rows,cw=[6,3,3,3])
    doc.add_paragraph()

    # === 六、诊断 ===
    add_h(doc,"六、数据诊断与建议",level=1)
    for ccy in ccy_order:
        t2=totals[ccy]; net=t2["net"]
        note=f"[{'流出' if net<0 else '流入'}] {ccy}：{'净流出'+fmt(abs(net),ccy) if net<0 else '净流入'+fmt(net,ccy)}"
        if net<0: note+=f"，占该币种支出 {pct(abs(net),t2['expense'])}"
        p=doc.add_paragraph(); r=p.add_run(note); r.bold=True
        tr=by_ccy[ccy].get("集团内收付",{})
        if tr.get("count",0)>0:
            tfp=(tr["income"]+tr["expense"])/max(t2["income"]+t2["expense"],1)*100
            doc.add_paragraph(f"  集团内收付占比 {tfp:.1f}%")
        doc.add_paragraph()

    doc.save(output_path)
    print(f"[完成] 报告 -> {output_path}")


# ============================================================
# 6. 明细导出
# ============================================================
def export_detail_xlsx(txns, output_dir, month_str):
    import xlsxwriter
    path=os.path.join(output_dir,f"资金明细_{month_str}.xlsx")
    wb=xlsxwriter.Workbook(path)
    for sn,inc in [("收入明细",True),("支出明细",False)]:
        ws=wb.add_worksheet(sn)
        hd=["日期","交易流水号","公司","币种","金额","账户余额","对方户名","本方开户行","本方账号",
           "对方开户行","对方账号","备注","交易备注","大类","子类"]
        for ci,h in enumerate(hd): ws.write(0,ci,h)
        r=1
        for t in txns:
            amt=t.income if inc else t.expense
            if amt<=0: continue
            a,s=classify(t,is_income=inc)
            ws.write(r,0,t.date.strftime("%Y-%m-%d")); ws.write(r,1,t.trade_no)
            ws.write(r,2,t.company); ws.write(r,3,t.currency)
            ws.write(r,4,round(amt,2)); ws.write(r,5,round(t.balance,2))
            ws.write(r,6,t.counterparty); ws.write(r,7,t.bank_self)
            ws.write(r,8,t.account_self); ws.write(r,9,t.bank_counter)
            ws.write(r,10,t.account_counter); ws.write(r,11,t.note)
            ws.write(r,12,t.trade_note); ws.write(r,13,a); ws.write(r,14,s)
            r+=1
    wb.close()
    print(f"[完成] 明细 -> {path}")


# ============================================================
# 7. 主入口
# ============================================================
def main():
    BASE_DIR=get_base_dir()
    ap=argparse.ArgumentParser(description="集团资金月度流量分析报告 — 自动生成")
    ap.add_argument("--input","-i",default=None,help="台账路径")
    ap.add_argument("--month","-m",default=None,help="月份 YYYY-MM（配合 --input 使用）")
    ap.add_argument("--output","-o",default=None,help="输出路径")
    args=ap.parse_args()

    if args.input:
        # 指定了文件路径
        input_path = args.input
        if not os.path.exists(input_path):
            print(f"[错误] 找不到文件: {input_path}"); return
        if args.month:
            ms = args.month
        else:
            ms = parse_month_from_filename(os.path.basename(input_path))
            if not ms:
                print(f"[错误] 无法从文件名推断月份，请使用 --month 指定"); return
    else:
        # 扫描目录，让用户选择
        all_files = scan_input_files(BASE_DIR)
        if not all_files:
            print(f"[错误] 目录中没有找到 资金台账*.xlsx 文件"); print(f"       当前: {BASE_DIR}"); return
        input_path, ms = select_file_interactive(all_files)

    print(f"[读取] {input_path}"); print(f"[月份] {ms}")
    txns=load_transactions(input_path,ms)
    rates=load_exchange_rates(input_path)
    print(f"[数据] {len(txns)} 笔")
    print(f"[汇率] {', '.join(f'{k}={v}' for k,v in rates.items() if k!='人民币')}")
    if not txns: print("[提示] 无数据"); return

    stats=analyze(txns,exchange_rates=rates)
    rmb_inc=stats.get("rmb_total_inc",0)/10000
    rmb_exp=stats.get("rmb_total_exp",0)/10000
    print(f"[折RMB] 收入:{rmb_inc:,.2f}万元  支出:{rmb_exp:,.2f}万元  净:{rmb_inc-rmb_exp:,.2f}万元")
    print(f"\n{'='*50}")
    for ccy in stats["ccy_order"]:
        t2=stats["totals"][ccy]
        print(f"[{ccy}] 收入:{fmt(t2['income'],ccy):>15s}  支出:{fmt(t2['expense'],ccy):>15s}  净:{fmt(t2['net'],ccy):>15s}")

    op=args.output or os.path.join(BASE_DIR,f"资金流量报告_{ms}_by_zion.docx")
    build_report(txns,ms,op,stats)
    export_detail_xlsx(txns,BASE_DIR,ms)
    print(f"\n[完成] OK！")

if __name__=="__main__":
    main()